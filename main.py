import os
from datetime import datetime
from pathlib import Path
from reports import init_reports_db, handle_report_message, start_weekly_reports

import requests
from flask import Flask, request
from docx import Document

app = Flask(__name__)

BOT_TOKEN = os.getenv("BOT_TOKEN")
TELEGRAM_API = f"https://api.telegram.org/bot{BOT_TOKEN}"

BASE_DIR = Path(__file__).resolve().parent
MOSCOW_CARGO_TEMPLATE_PATH = BASE_DIR / "templates" / "moscow_cargo_template.docx"
PORTS_RF_TEMPLATE_PATH = BASE_DIR / "templates" / "ports_rf_template.docx"
GENERATED_DIR = BASE_DIR / "generated"
GENERATED_DIR.mkdir(exist_ok=True)

user_states = {}
power_counter = 0


def send_message(chat_id, text, reply_markup=None):
    payload = {
        "chat_id": chat_id,
        "text": text
    }

    if reply_markup:
        payload["reply_markup"] = reply_markup

    requests.post(f"{TELEGRAM_API}/sendMessage", json=payload)


def send_document(chat_id, file_path, caption=""):
    with open(file_path, "rb") as file:
        requests.post(
            f"{TELEGRAM_API}/sendDocument",
            data={
                "chat_id": chat_id,
                "caption": caption
            },
            files={
                "document": file
            },
        )


def send_main_menu(chat_id):
    keyboard = {
        "inline_keyboard": [
            [
                {
                    "text": "📄 Доверенность Москва Карго",
                    "callback_data": "power_moscow_cargo"
                }
            ],
            [
                {
                    "text": "📄 Универсальная доверенность (Порты РФ)",
                    "callback_data": "power_ports_rf"
                }
            ],
        ]
    }

    send_message(chat_id, "Выберите действие:", keyboard)


@app.route("/", methods=["GET"])
def index():
    return "Bot is running"


@app.route("/webhook", methods=["POST"])
def webhook():
    update = request.get_json()

    if "message" in update:
        handle_message(update["message"])

    if "callback_query" in update:
        handle_callback(update["callback_query"])

    return "ok"


def handle_message(message):
    chat_id = message["chat"]["id"]
    chat_type = message["chat"]["type"]
    text = message.get("text", "").strip()

    # В группах бот принимает только ежедневные сводки
    if chat_type != "private":
        handle_report_message(message, send_message)
        return

    if text == "/start":
        send_main_menu(chat_id)
        return

    process_power_input(chat_id, text)


def handle_callback(callback):
    chat_id = callback["message"]["chat"]["id"]
    data = callback["data"]

    requests.post(
        f"{TELEGRAM_API}/answerCallbackQuery",
        json={"callback_query_id": callback["id"]},
    )

    if data == "power_moscow_cargo":
        user_states[chat_id] = {
            "step": "company",
            "template_type": "moscow_cargo",
            "data": {}
        }

        send_message(chat_id, "Введите название компании-доверителя:")
        return

    if data == "power_ports_rf":
        user_states[chat_id] = {
            "step": "city",
            "template_type": "ports_rf",
            "data": {}
        }

        send_message(
            chat_id,
            "Введите место выдачи доверенности.\nНапример: г. Новосибирск, Российская Федерация"
        )
        return

    if data in ["transfer_yes", "transfer_no"]:
        state = user_states.get(chat_id)

        if not state:
            return

        if data == "transfer_yes":
            state["data"]["transfer_right"] = "с правом последующего передоверия"
        else:
            state["data"]["transfer_right"] = "без права передоверия"

        state["step"] = "term"
        send_message(chat_id, "Введите срок действия доверенности. Например: 1 год")
        return


def process_power_input(chat_id, text):
    state = user_states.get(chat_id)

    if not state:
        return

    step = state["step"]
    data = state["data"]

    if step == "city":
        data["city"] = text
        state["step"] = "company"
        send_message(chat_id, "Введите название компании-доверителя:")
        return

    if step == "company":
        data["company"] = text
        state["step"] = "address"
        send_message(chat_id, "Введите адрес компании:")
        return

    if step == "address":
        data["address"] = text
        state["step"] = "inn"
        send_message(chat_id, "Введите ИНН:")
        return

    if step == "inn":
        data["inn"] = text
        state["step"] = "ogrn"
        send_message(chat_id, "Введите ОГРН:")
        return

    if step == "ogrn":
        data["ogrn"] = text
        state["step"] = "director_position"
        send_message(chat_id, "Введите должность руководителя. Например: генерального директора")
        return

    if step == "director_position":
        data["director_position"] = text
        state["step"] = "director_name"
        send_message(
            chat_id,
            "Введите ФИО руководителя в родительном падеже. Например: Иванова Ивана Ивановича"
        )
        return

    if step == "director_name":
        data["director_name"] = text
        state["step"] = "transfer_right"

        keyboard = {
            "inline_keyboard": [
                [
                    {
                        "text": "Да, с правом передоверия",
                        "callback_data": "transfer_yes"
                    }
                ],
                [
                    {
                        "text": "Нет, без права передоверия",
                        "callback_data": "transfer_no"
                    }
                ],
            ]
        }

        send_message(chat_id, "Доверенность с правом передоверия?", keyboard)
        return

    if step == "term":
        data["term"] = text
        template_type = state.get("template_type", "moscow_cargo")

        user_states.pop(chat_id, None)

        send_message(chat_id, "Формирую доверенность...")

        try:
            docx_path = generate_power_document(data, template_type)
            send_document(chat_id, docx_path, "Готово: Word-файл доверенности")
        except Exception as error:
            send_message(chat_id, f"❌ Ошибка при формировании файла:\n{error}")
            print(f"ERROR GENERATING DOCUMENT: {error}", flush=True)

        send_main_menu(chat_id)
        return


def generate_power_document(data, template_type):
    number = generate_number()
    date_text = datetime.now().strftime("%d.%m.%Y")

    replacements = {
        "{{NUMBER}}": number,
        "{{DATE}}": date_text,
        "{{COMPANY}}": data["company"],
        "{{ADDRESS}}": data["address"],
        "{{INN}}": data["inn"],
        "{{OGRN}}": data["ogrn"],
        "{{DIRECTOR_POSITION}}": data["director_position"],
        "{{DIRECTOR_NAME}}": data["director_name"],
        "{{TERM}}": data["term"],
        "{{TRANSFER_RIGHT}}": data["transfer_right"],
        "{{CITY}}": data.get("city", "г. Москва"),
        "{{DATE_TEXT}}": get_russian_date_text(),
    }

    if template_type == "ports_rf":
        template_path = PORTS_RF_TEMPLATE_PATH
        file_prefix = "Доверенность_Порты_РФ"
    else:
        template_path = MOSCOW_CARGO_TEMPLATE_PATH
        file_prefix = "Доверенность_Москва_Карго"

    if not template_path.exists():
        raise FileNotFoundError(f"Не найден шаблон: {template_path}")

    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

    safe_company = make_safe_filename(data["company"])
    filename = f"{file_prefix}_{number}_{safe_company}.docx"
    output_path = GENERATED_DIR / filename

    doc.save(output_path)
    return output_path


def replace_in_paragraph(paragraph, replacements):
    full_text = "".join(run.text for run in paragraph.runs)

    if not any(key in full_text for key in replacements):
        return

    for key, value in replacements.items():
        full_text = full_text.replace(key, str(value))

    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        paragraph.runs[0].text = full_text
    else:
        paragraph.add_run(full_text)


def generate_number():
    global power_counter

    power_counter += 1
    year = datetime.now().strftime("%Y")

    return f"ДМК-{year}-{str(power_counter).zfill(3)}"


def make_safe_filename(text):
    bad_chars = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']

    for char in bad_chars:
        text = text.replace(char, "_")

    return text.replace(" ", "_")[:60]


def get_russian_date_text():
    now = datetime.now()

    days = {
        1: "первое",
        2: "второе",
        3: "третье",
        4: "четвертое",
        5: "пятое",
        6: "шестое",
        7: "седьмое",
        8: "восьмое",
        9: "девятое",
        10: "десятое",
        11: "одиннадцатое",
        12: "двенадцатое",
        13: "тринадцатое",
        14: "четырнадцатое",
        15: "пятнадцатое",
        16: "шестнадцатое",
        17: "семнадцатое",
        18: "восемнадцатое",
        19: "девятнадцатое",
        20: "двадцатое",
        21: "двадцать первое",
        22: "двадцать второе",
        23: "двадцать третье",
        24: "двадцать четвертое",
        25: "двадцать пятое",
        26: "двадцать шестое",
        27: "двадцать седьмое",
        28: "двадцать восьмое",
        29: "двадцать девятое",
        30: "тридцатое",
        31: "тридцать первое",
    }

    months = {
        1: "января",
        2: "февраля",
        3: "марта",
        4: "апреля",
        5: "мая",
        6: "июня",
        7: "июля",
        8: "августа",
        9: "сентября",
        10: "октября",
        11: "ноября",
        12: "декабря",
    }

    return f"{days[now.day]} {months[now.month]} две тысячи двадцать шестого года"


if __name__ == "__main__":
    init_reports_db()
    start_weekly_reports(send_message)
    app.run(host="0.0.0.0", port=5000)
